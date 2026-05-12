from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from glob import glob
import os 
import json
import re

def get_docx_txt(file_path):
    try:
        document = Document(file_path)
    except PackageNotFoundError as e:
        print(f"Package not found: {e}")
        return None, None
    except Exception as e:
        print(f"Package not found: {e}")
        return None, None

    docx_text = ""
    for para in document.paragraphs:
        docx_text += para.text + "\n"

    for table in document.tables:
        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ''
                for para in cell.paragraphs:
                    cell_text += para.text + '\n'
                cell_text = cell_text.strip()
                
                cell_text = cell_text.replace('\n', ' ')
                cell_text = cell_text.replace('\|', ' ')
                cell_text = cell_text.replace('|', ' ')

                row_data.append(cell_text)
            rows.append(row_data)
        
        header = rows[0]
        docx_text += "| " + " | ".join(header) + " |\n"
        docx_text += "| " + " | ".join(['---'] * len(header)) + " |\n"

        for row in rows[1:]:
            docx_text += "| " + " | ".join(row) + " |\n"

        docx_text += "\n"  


    for content in document.paragraphs:
        if content.style.name=='Heading 1' or content.style.name=='Heading 2' or content.style.name=='Heading 3':
            docx_text += content.text + "\n"
    
    return docx_text, document.core_properties

def docx_parser(docx_folder, temp_txt_folder, method="fast"):
    types = ('*.docx', '*.doc')
    articleFile = []
    for files in types:
        articleFile.extend(glob(docx_folder + "/" + files))
    
    titleJson = {}
    output_folder = os.path.join(temp_txt_folder, "docs_txt")
    os.makedirs(output_folder, exist_ok=True)  

    if not os.path.exists(temp_txt_folder + "/titleJson.json"):
        titleJson = {}
    else:
        with open(temp_txt_folder + "/titleJson.json", "r", encoding="utf-8") as fp:
            titleJson = json.load(fp)
    

    for af in articleFile:
        file_name = os.path.splitext(af.split("/")[-1])[0]
        txt_file = os.path.join(output_folder, f"{file_name}.txt")  

        if method == "fast":
            ori_txtData, metadata = get_docx_txt(af)

        if ori_txtData is not None:
            with open(txt_file, "w", encoding="utf-8") as fp:
                fp.write(ori_txtData)
            
            titleJson[file_name] = af.split("/")[-1]
    
    with open(temp_txt_folder + "/titleJson.json", "w", encoding="utf-8") as fp:
        json.dump(titleJson, fp, indent=4, ensure_ascii=False)
